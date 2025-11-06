import logging
import io
from typing import Dict, Any, Optional
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentParser:
    """文档解析器，支持多种文档格式"""
    
    @staticmethod
    def parse(file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        根据文件扩展名自动解析文档
        
        Args:
            file_content: 文件二进制内容
            filename: 文件名（用于判断格式）
            
        Returns:
            包含解析结果的字典，格式: {"text": str, "metadata": dict}
        """
        file_ext = Path(filename).suffix.lower()
        
        try:
            if file_ext in ['.docx', '.doc']:
                return DocumentParser._parse_word(file_content)
            elif file_ext == '.pdf':
                return DocumentParser._parse_pdf(file_content)
            elif file_ext in ['.xlsx', '.xls']:
                return DocumentParser._parse_excel(file_content)
            elif file_ext == '.xmind':
                return DocumentParser._parse_xmind(file_content)
            else:
                raise ValueError(f"不支持的文件格式: {file_ext}")
        except Exception as e:
            logger.error(f"解析文件 {filename} 失败: {e}")
            raise ValueError(f"文件解析失败: {str(e)}")
    
    @staticmethod
    def _parse_word(file_content: bytes) -> Dict[str, Any]:
        """解析Word文档"""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            
            return {
                "text": text,
                "metadata": {
                    "format": "word",
                    "paragraphs_count": len(doc.paragraphs),
                    "tables_count": len(doc.tables)
                }
            }
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")
    
    @staticmethod
    def _parse_pdf(file_content: bytes) -> Dict[str, Any]:
        """解析PDF文档"""
        try:
            import pdfplumber
            
            text_parts = []
            total_pages = 0
            
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                total_pages = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text.strip())
            
            text = "\n\n".join(text_parts)
            
            return {
                "text": text,
                "metadata": {
                    "format": "pdf",
                    "pages_count": total_pages
                }
            }
        except ImportError:
            # 回退到PyPDF2
            try:
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                text_parts = []
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text.strip())
                
                text = "\n\n".join(text_parts)
                
                return {
                    "text": text,
                    "metadata": {
                        "format": "pdf",
                        "pages_count": len(pdf_reader.pages)
                    }
                }
            except ImportError:
                raise ImportError("请安装 pdfplumber 或 PyPDF2: pip install pdfplumber PyPDF2")
    
    @staticmethod
    def _parse_excel(file_content: bytes) -> Dict[str, Any]:
        """解析Excel文档"""
        try:
            import openpyxl
            
            workbook = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True)
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"=== 工作表: {sheet_name} ===")
                
                # 提取所有单元格的文本
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell_value in row:
                        if cell_value is not None:
                            row_text.append(str(cell_value).strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
                text_parts.append("")  # 工作表之间添加空行
            
            text = "\n".join(text_parts)
            
            return {
                "text": text,
                "metadata": {
                    "format": "excel",
                    "sheets_count": len(workbook.sheetnames),
                    "sheet_names": workbook.sheetnames
                }
            }
        except ImportError:
            # 回退到xlrd（用于.xls格式）
            try:
                import xlrd
                workbook = xlrd.open_workbook(file_contents=file_content)
                text_parts = []
                
                for sheet_name in workbook.sheet_names():
                    sheet = workbook.sheet_by_name(sheet_name)
                    text_parts.append(f"=== 工作表: {sheet_name} ===")
                    
                    for row_idx in range(sheet.nrows):
                        row = sheet.row_values(row_idx)
                        row_text = [str(cell).strip() for cell in row if cell]
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                    text_parts.append("")
                
                text = "\n".join(text_parts)
                
                return {
                    "text": text,
                    "metadata": {
                        "format": "excel",
                        "sheets_count": len(workbook.sheet_names()),
                        "sheet_names": workbook.sheet_names()
                    }
                }
            except ImportError:
                raise ImportError("请安装 openpyxl 或 xlrd: pip install openpyxl xlrd")
    
    @staticmethod
    def _parse_xmind(file_content: bytes) -> Dict[str, Any]:
        """解析XMind思维导图"""
        try:
            import xmindparser
            import tempfile
            import os
            
            # XMind文件需要保存为临时文件来解析
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xmind') as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name
            
            try:
                # 解析XMind文件
                xmind_content = xmindparser.xmind_to_dict(tmp_file_path)
                
                def extract_text_from_topic(topic: dict, level: int = 0) -> list:
                    """递归提取思维导图文本"""
                    text_parts = []
                    indent = "  " * level
                    
                    if isinstance(topic, dict):
                        title = topic.get('title', '')
                        if title and title.strip():
                            text_parts.append(f"{indent}{title}")
                        
                        # 处理子主题
                        children = topic.get('topics', [])
                        if children and isinstance(children, list):
                            for child in children:
                                if isinstance(child, dict):
                                    text_parts.extend(extract_text_from_topic(child, level + 1))
                                elif isinstance(child, list):
                                    # 有时topics可能是嵌套列表
                                    for sub_child in child:
                                        if isinstance(sub_child, dict):
                                            text_parts.extend(extract_text_from_topic(sub_child, level + 1))
                    elif isinstance(topic, list):
                        # 如果topic本身是列表，遍历处理
                        for item in topic:
                            if isinstance(item, dict):
                                text_parts.extend(extract_text_from_topic(item, level))
                    
                    return text_parts
                
                text_parts = []
                # XMind结构可能是多种格式
                if isinstance(xmind_content, dict):
                    # 处理 {"xmindcontent": [...]} 格式
                    for key, value in xmind_content.items():
                        if isinstance(value, list):
                            for item in value:
                                text_parts.extend(extract_text_from_topic(item))
                        elif isinstance(value, dict):
                            text_parts.extend(extract_text_from_topic(value))
                elif isinstance(xmind_content, list):
                    # 直接是列表格式
                    for item in xmind_content:
                        text_parts.extend(extract_text_from_topic(item))
                
                text = "\n".join(text_parts) if text_parts else ""
                
                if not text.strip():
                    # 如果提取不到文本，尝试从原始内容中提取
                    logger.warning("未能从XMind中提取到文本，尝试其他方法")
                    # 可以尝试使用xmindparser的其他方法
                    try:
                        # xmindparser可能还有其他解析方法
                        workbook = xmindparser.xmind_to_workbook(tmp_file_path)
                        if workbook:
                            # 尝试从workbook中提取
                            for sheet in workbook.getSheets():
                                root_topic = sheet.getRootTopic()
                                if root_topic:
                                    text_parts.append(root_topic.getTitle() or "")
                                    for topic in root_topic.getSubTopics():
                                        text_parts.append(f"  {topic.getTitle() or ''}")
                            text = "\n".join(text_parts) if text_parts else ""
                    except:
                        pass
                
                return {
                    "text": text,
                    "metadata": {
                        "format": "xmind",
                        "text_extracted": len(text) > 0
                    }
                }
            finally:
                # 删除临时文件
                if os.path.exists(tmp_file_path):
                    try:
                        os.unlink(tmp_file_path)
                    except:
                        pass
        except ImportError:
            raise ImportError("请安装 xmindparser: pip install xmindparser")
        except Exception as e:
            logger.error(f"解析XMind文件失败: {e}", exc_info=True)
            raise ValueError(f"XMind文件解析失败: {str(e)}")
